from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
  pyttsx3.speak(text)
document = Document()


document.add_picture(
    'sebastian.jpg',
    width=Inches(2.0)
)

name = input('What is your name?')
speak('Hello ' + name + ' how are you to day?')

speak('What is your phone number?')
phone_number = input('What is your phone number?')
email = input('What is your email?')
document.add_paragraph(

    name + ' | ' + phone_number + ' | ' + email
)

document.add_heading('About me')

document.add_paragraph(
input('Tell about your self')

)

document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From Date')
to_date = input('To Date ')

p.add_run(company + '').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True
experience_details = input(
    'Describe your experience at ' + company
)
p.add_run(experience_details)

while True:
  has_more_experience = input('Do you have more experiences? Yes or No ')
  if has_more_experience.lower() == 'yes':
    p = document.add_paragraph()
    company = input('Enter Company ')
    from_date = input('From Date ')
    to_date = input('To date ')
    p.add_run(company + ' ').bold = True
    p.add_run(from_date + '-' + to_date + '\n').italic = True
    experience_details = input(
    'Describe your experience at ' + company + ' '
    )
    p.add_run(experience_details)

  else:
    break

#skills
document.add_heading('Skills')
skill = input('Enter Skills ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

#more experience
while True:
  has_more_skill = input('Do you have more skills? Yes or No ')
  if has_more_skill.lower() == 'yes':
    skill = input('Enter skils')
    p = document.add_paragraph (skill)
    p.style = 'List Bullet'


    p.add_run(company + '').bold = True
    p.add_run(from_date + '-' + to_date + '\n').italic = True
    experience_details = input(
        'Describe your experience at ' + company + ' '
    )
    p.add_run(experience_details)
  else:
    break
  #footer
  section = document.sections[0]
  footer = section.footer
  p = footer.paragraph[0]
  p.text = "CV generated using Valentino"


document.save('cv.docx')
