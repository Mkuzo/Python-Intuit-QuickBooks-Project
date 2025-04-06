import pyttsx3
from docx import Document
from docx.shared import Inches

# Initialize text-to-speech
engine = pyttsx3.init()

def speak(text):
    engine.say(text)
    engine.runAndWait()

# Create a new document
document = Document()

# Add profile picture
try:
    document.add_picture('sebastian.jpg', width=Inches(2.0))
except:
    speak("Profile picture not found. Skipping image.")

# Basic Info
speak("What is your name?")
name = input('What is your name? ')

speak("What is your phone number?")
phone_number = input('What is your phone number? ')

speak("What is your email?")
email = input('What is your email? ')

document.add_paragraph(name + ' | ' + phone_number + ' | ' + email)

# About Me
document.add_heading('About Me')
speak("Tell me about yourself.")
about_me = input('Tell about yourself: ')
document.add_paragraph(about_me)

# Work Experience
document.add_heading('Work Experience')

while True:
    p = document.add_paragraph()
    
    speak("Enter company name.")
    company = input('Enter company: ')

    speak("From which date?")
    from_date = input('From Date: ')

    speak("To which date?")
    to_date = input('To Date: ')

    p.add_run(company + ' ').bold = True
    p.add_run(from_date + ' - ' + to_date + '\n').italic = True

    speak(f'Describe your experience at {company}.')
    experience_details = input(f'Describe your experience at {company}: ')
    p.add_run(experience_details)

    speak("Do you have more work experience? Say yes or no.")
    more = input('Do you have more experiences? (Yes or No): ')
    if more.lower() != 'yes':
        break

# Skills Section
document.add_heading('Skills')
speak("Let's add your skills. Say done when you are finished.")

while True:
    skill = input('Enter a skill (or type "done" to finish): ')
    if skill.lower() == 'done':
        break
    p = document.add_paragraph(skill)
    p.style = 'List Bullet'

# Save the document
document.save('Cv.docx')
speak("Your CV has been created and saved as cv.docx.")
